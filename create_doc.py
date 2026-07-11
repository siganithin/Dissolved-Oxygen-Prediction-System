import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 153)
        run.font.name = 'Calibri'
    return heading

def add_qa(doc, question, answer):
    p = doc.add_paragraph()
    q_run = p.add_run(f"Q: {question}\n")
    q_run.bold = True
    q_run.font.color.rgb = RGBColor(192, 0, 0)
    
    a_run = p.add_run(f"A: {answer}")
    a_run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_after = Pt(12)

def main():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Expert Viva Preparation Guide\nDissolved Oxygen Prediction System\n', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("This document is a deeply humanized, expert-level preparation guide designed specifically for your Viva Voice. It breaks down technical concepts into easily understandable insights, perfect for answering tricky examiner questions.\n")

    # Section 1: Dashboard & Results
    add_heading(doc, "1. Deep Dive: The Dashboard & Visual Results", 1)
    
    add_qa(doc, "What tools did you use to build the dashboard, and why only Streamlit?", 
           "I used Streamlit entirely in Python. The reason I strictly chose Streamlit (instead of React, Angular, or Django) is because our core project is heavily Data Science and Machine Learning focused. Streamlit is natively integrated with Pandas, NumPy, and Keras. It allows us to deploy interactive machine learning predictive apps directly from our Python scripts without writing heavy HTML/CSS frontend logic. It provided the exact components we needed: dynamic sliders, metric cards, and instant Matplotlib chart rendering.")
    
    add_qa(doc, "How does the threshold slider actually work under the hood?",
           "The threshold slider captures a user-defined float value (defaulting at 5.0 mg/L). Every time the model makes a DO prediction, the predicted value is compared against this slider's value in real-time. If the prediction drops below the threshold, a pandas DataFrame column 'alert' gets tagged as 'True'. Streamlit instantly reads this DataFrame and conditionally renders warning alerts and adjusts the metric cards.")
           
    add_qa(doc, "How did you optimize the dashboard so it doesn't freeze or slow down?",
           "Streamlit has a quirk where it re-runs the entire python script every time a slider is moved. This means our heavy deep learning model would reload and predict data thousands of times! To fix this, I engineered a caching mechanism using the @st.cache_data decorator. Now, when a CSV is uploaded, the predictions are calculated precisely once and saved in local memory. Any changes to the UI slider instantly use the cached data, making the system run blazingly fast.")

    # Section 2: Diagrams & System Flow
    add_heading(doc, "2. Explaining Project UML Diagrams", 1)
    
    add_qa(doc, "How would you explain your Use Case Diagram?",
           "The primary 'Actor' in our Use Case Diagram is the Aquaculturist (the farmer). Their main use cases are uploading sensor data, monitoring trends, and receiving alerts. Under the hood, these actions 'include' (via <<includes>> relationships) system sub-processes like Data Preprocessing and Feature Selection, which happen automatically without the farmer's intervention.")
           
    add_qa(doc, "What does the Activity Diagram represent in your project?",
           "The Activity diagram tracks the real-world flow of operations. It starts with raw Data Collection -> passes to Preprocessing (handling nulls/scaling) -> routes to LightGBM for Top-Feature extraction -> enters the BiSRU-Attention neural network -> and finally hits a decision node: 'Is DO < Threshold?'. If yes, it branches to 'Trigger Alert'; if no, it loops back to normal monitoring.")

    add_qa(doc, "What is the architectural flow of your system?",
           "Our System Architecture is a strict sequential pipeline: Sensor Data -> Cleaning -> Feature Extraction (LightGBM) -> Sequence Modeling (BiSRU-Attention) -> Output Dashboard (Streamlit). Each layer acts as an isolated pipeline stage, allowing us to easily swap out models in the future without breaking the UI.")

    # Section 3: Machine Learning & Deep Learning Basics
    add_heading(doc, "3. Core Machine Learning & Deep Learning Concepts", 1)

    add_qa(doc, "What exactly is Data Preprocessing, and why did you focus heavily on it?",
           "Data gathered in the real world is inherently 'dirty'—it has missing slots, negative erroneous values, or huge variations in scales. If we feed this raw data into an ML model, it learns 'garbage in, garbage out'. Preprocessing is cleaning it. Specifically, we used column medians to fill missing values (so we don't skew the distribution) and dropped negative rows.")
           
    add_qa(doc, "What is Min-Max Normalization and why is it crucial?",
           "Min-Max Normalization mathematically squashes all features into a range between 0 and 1. Here is the logic: Temperature might be 30, but Ammonia might be 0.02. To a neural network, 30 is heavily prioritized over 0.02 just because it's a larger number! Normalizing guarantees the network treats every parameter equally, speeding up the 'Gradient Descent' convergence so the network learns faster and without bias.")
           
    add_qa(doc, "What is the mathematical difference between Machine Learning and Deep Learning in your project context?",
           "Machine Learning (like our LightGBM) relies on human-crafted or structured statistical rules (like Decision Trees) to map inputs to outputs. It's fast but struggles with complex chronological sequences. Deep Learning (our BiSRU network) uses interconnected nodes that mimic the human brain (Neural Networks). It autonomously learns hidden temporal relationships over time sequences (the past 24 hours of water quality), which traditional ML cannot do efficiently.")

    add_qa(doc, "Why Deep Learning for Time Series?",
           "Because traditional ML models treat every row of data as independent. But in nature, the oxygen level at 2:00 PM is heavily dependent on the oxygen level at 1:00 PM. Recurrent Neural Networks (like our BiSRU) contain 'memory'—they pass the output of the past state into the current state as input, explicitly retaining chronological history.")

    # Section 4: Performance Evaluation Metrics
    add_heading(doc, "4. Evaluation Metrics & Mathematical Breakdown", 1)
    
    doc.add_paragraph("Evaluators love asking about the math behind accuracy. Here is how our model was measured:\n")

    # MSE
    p1 = doc.add_paragraph()
    p1.add_run("Mean Squared Error (MSE): ").bold = True
    p1.add_run("Our model achieved 0.0008.\n")
    p1.add_run("Formula: MSE = (1/n) * Σ(Actual - Predicted)²\n").italic = True
    p1.add_run("Explanation: It measures the average of the squares of the errors. We square it for two reasons: a) ensuring negative errors don't cancel out positive errors, and b) heavily punishing large outliers. An MSE near 0 is outstanding.")

    # MAE
    p2 = doc.add_paragraph()
    p2.add_run("Mean Absolute Error (MAE): ").bold = True
    p2.add_run("Our model achieved 0.0199.\n")
    p2.add_run("Formula: MAE = (1/n) * Σ|Actual - Predicted|\n").italic = True
    p2.add_run("Explanation: It calculates the absolute average distance between predictions and actuals. Unlike MSE, it treats all errors equally without exponentially punishing outliers. It tells us that our DO predictions are, on average, incredibly close (off by only 0.0199) to reality.")

    # RMSE
    p3 = doc.add_paragraph()
    p3.add_run("Root Mean Squared Error (RMSE): ").bold = True
    p3.add_run("Our model achieved 0.0285.\n")
    p3.add_run("Formula: RMSE = √ [ (1/n) * Σ(Actual - Predicted)² ]\n").italic = True
    p3.add_run("Explanation: It is simply the square root of the MSE. We use RMSE because MSE's units are squared (e.g., (mg/L)²), which doesn't make logical sense when reporting. RMSE brings the error metric back to the original unit base (mg/L).")

    # R2
    p4 = doc.add_paragraph()
    p4.add_run("R-Squared (R² / Coefficient of Determination): ").bold = True
    p4.add_run("Our model achieved 96.28% (0.9628).\n")
    p4.add_run("Formula: R² = 1 - (SSR / SST)\n").italic = True
    p4.add_run("Explanation: Where SSR is the sum of squared residuals and SST is the total sum of squares. In human terms: It represents the percentage of variance in Dissolved Oxygen that our model completely explains. 96.28% is phenomenal and proves our selection of LightGBM features and Attention mechanism perfectly capture the volatility in the water.")

    # Section 5: Bonus Expert Level Questions
    add_heading(doc, "5. The Secret Sauce (Pro-Level Viva Answers)", 1)
    
    add_qa(doc, "What exactly does the 'Attention Layer' do mathematically?",
           "Standard neural networks force all 24 past hours of data into one summarized vector block, which causes data loss. The Attention layer calculates a mathematical 'Score' (using a Tanh activation function) for every single historical hour. It then converts these scores to probabilities (Weights) using a Softmax function. It multiplies the data by these weights—essentially saying 'Ignore hour 3, but zoom in on hour 19 because the ammonia spiked'.")
           
    add_qa(doc, "Did your system have any overfitting? How did you control it?",
           "We strictly mitigated overfitting using 'Early Stopping' and a 'Validation Split' (15%). During training, the model's performance on unseen validation data was checked every epoch. If the model started memorizing the training data instead of generalizing (meaning training loss went down, but validation loss went up), the Early Stopping mechanism forcibly halted the training process at the 5-epoch patience limit, actively restoring the best mathematical weights.")

    doc.add_paragraph("\n*** Best of luck with your Viva Voice! You've deeply understood and implemented highly advanced data science principles. Defend with confidence! ***")
    
    doc.save("EXPERT_VIVA_PREPARATION_GUIDE.docx")

if __name__ == "__main__":
    main()
